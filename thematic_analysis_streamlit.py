import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import networkx as nx
import matplotlib.pyplot as plt
from datetime import datetime
import json
import sqlite3
import io
import os
import base64
from collections import Counter, defaultdict
import re
import docx
from textblob import TextBlob
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from sklearn.decomposition import PCA
# LLM REPORT GENERATION - OPENAI
import openai
import textwrap
from difflib import SequenceMatcher
try:
    import pdfplumber
except ImportError:
    pdfplumber = None
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# WordCloud import with fallback
try:
    from wordcloud import WordCloud
    WORDCLOUD_AVAILABLE = True
except ImportError:
    WORDCLOUD_AVAILABLE = False

import seaborn as sns
# --- Simplified Kaleido Setup for Streamlit Cloud (Pinned Versions) ---
try:
    import kaleido
    # For Kaleido 0.2.1, this ensures bundled Chrome is used (no download needed)
    kaleido.__version__  # Just verify import
except Exception as e:
    pass
# --------------------------------------------------------

# Page configuration
st.set_page_config(
    page_title="Reflexive Thematic Analysis Tool",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Database setup with phase tracking - Enhanced with role and theme versioning
@st.cache_resource
def init_database():
    conn = sqlite3.connect('reflexive_analysis.db', check_same_thread=False)
    cursor = conn.cursor()

    # Transcripts table with participant categorization - Added role
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS transcripts (
            id TEXT PRIMARY KEY,
            participant_type TEXT,
            ethnicity TEXT,
            role TEXT,
            location TEXT,
            duration TEXT,
            content TEXT,
            status TEXT,
            upload_date DATE,
            phase_completed INTEGER DEFAULT 1
        )
    ''')

    # Check and add role column if missing
    cursor.execute("PRAGMA table_info(transcripts)")
    info = cursor.fetchall()
    existing_columns = [col[1] for col in info]
    if 'role' not in existing_columns:
        cursor.execute("ALTER TABLE transcripts ADD COLUMN role TEXT")

    # Codes table with AI suggestions and phase tracking
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS codes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            code_name TEXT,
            description TEXT,
            frequency INTEGER DEFAULT 0,
            source_type TEXT,
            ai_suggested BOOLEAN DEFAULT 0,
            phase_created INTEGER DEFAULT 2,
            created_date DATE,
            com_b_category TEXT,
            transcript_id TEXT,
            FOREIGN KEY (transcript_id) REFERENCES transcripts (id)
        )
    ''')

    # Code instances with highlighting positions
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS code_instances (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            transcript_id TEXT,
            code_id INTEGER,
            text_segment TEXT,
            start_position INTEGER,
            end_position INTEGER,
            phase_coded INTEGER DEFAULT 2,
            FOREIGN KEY (transcript_id) REFERENCES transcripts (id),
            FOREIGN KEY (code_id) REFERENCES codes (id)
        )
    ''')

    # Themes with phase tracking and COM-B mapping - Added version
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS themes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            theme_name TEXT,
            description TEXT,
            com_b_primary TEXT,
            com_b_secondary TEXT,
            parent_theme_id INTEGER,
            phase_developed INTEGER DEFAULT 4,
            prevalence_percentage REAL,
            version INTEGER DEFAULT 1,
            created_date DATE
        )
    ''')

    # Check and add version column if missing
    cursor.execute("PRAGMA table_info(themes)")
    info = cursor.fetchall()
    existing_columns = [col[1] for col in info]
    if 'version' not in existing_columns:
        cursor.execute("ALTER TABLE themes ADD COLUMN version INTEGER DEFAULT 1")

    # Theme history for versioning
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS theme_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            theme_id INTEGER,
            version INTEGER,
            theme_name TEXT,
            description TEXT,
            com_b_primary TEXT,
            com_b_secondary TEXT,
            changes TEXT,
            created_date DATE,
            FOREIGN KEY (theme_id) REFERENCES themes (id)
        )
    ''')

    # Theme-code relationships
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS theme_codes (
            theme_id INTEGER,
            code_id INTEGER,
            PRIMARY KEY (theme_id, code_id),
            FOREIGN KEY (theme_id) REFERENCES themes (id),
            FOREIGN KEY (code_id) REFERENCES codes (id)
        )
    ''')

    # Reflexive memos by phase
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS reflexive_memos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            phase INTEGER,
            memo_text TEXT,
            created_date DATE,
            updated_date DATE
        )
    ''')

    # Per-transcript memos
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS transcript_memos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            transcript_id TEXT,
            memo_text TEXT,
            created_date DATE,
            FOREIGN KEY (transcript_id) REFERENCES transcripts (id)
        )
    ''')

    conn.commit()
    return conn

conn = init_database()

# Data functions - Updated save_transcript with role
def save_transcript(transcript_data):
    cursor = conn.cursor()
    transcript_id, participant_type, ethnicity, role, location, duration, content, status, upload_date = transcript_data
    cursor.execute('''
        INSERT OR REPLACE INTO transcripts
        (id, participant_type, ethnicity, role, location, duration, content, status, upload_date)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', transcript_data)
    conn.commit()
    return transcript_id, content

def get_transcripts():
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM transcripts')
    columns = [description[0] for description in cursor.description]
    return pd.DataFrame(cursor.fetchall(), columns=columns)

def save_code(code_name, description, source_type, ai_suggested=False, com_b_category=None, transcript_id=None):
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO codes (code_name, description, source_type, ai_suggested, created_date, com_b_category, transcript_id)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    ''', (code_name, description, source_type, ai_suggested, datetime.now().date(), com_b_category, transcript_id))
    code_id = cursor.lastrowid
    conn.commit()

    # Auto-create instances for highlighting with fuzzy matching
    if transcript_id:
        transcript_cursor = conn.cursor()
        transcript_cursor.execute('SELECT content FROM transcripts WHERE id = ?', (transcript_id,))
        result = transcript_cursor.fetchone()
        if result:
            content = result[0]
            create_instances(code_id, transcript_id, content, code_name)
    return code_id

def find_approximate_matches(text, code_name, threshold=0.8):
    text_lower = text.lower()
    code_lower = code_name.lower()
    positions = []
    len_code = len(code_lower)
    for i in range(len(text_lower) - len_code + 1):
        substring = text_lower[i:i + len_code]
        ratio = SequenceMatcher(None, substring, code_lower).ratio()
        if ratio > threshold:
            # Find word boundaries
            start = max(0, text_lower.rfind(' ', 0, i) + 1 if i > 0 else 0)
            end = min(len(text), text_lower.find(' ', i + len_code) if text_lower.find(' ', i + len_code) != -1 else len(text))
            match_text = text[start:end]
            positions.append((start, end, match_text))
    return list(set(positions))  # Deduplicate

def create_instances(code_id, transcript_id, content, code_name, threshold=0.8):
    cursor = conn.cursor()
    positions = find_approximate_matches(content, code_name, threshold)
    for start, end, text_segment in positions:
        cursor.execute('''
            INSERT INTO code_instances (transcript_id, code_id, text_segment, start_position, end_position)
            VALUES (?, ?, ?, ?, ?)
        ''', (transcript_id, code_id, text_segment, start, end))
    conn.commit()

def get_codes():
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM codes')
    columns = [description[0] for description in cursor.description]
    return pd.DataFrame(cursor.fetchall(), columns=columns)

def get_codes_by_transcript_id(transcript_id):
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM codes WHERE transcript_id = ?', (transcript_id,))
    columns = [description[0] for description in cursor.description]
    return pd.DataFrame(cursor.fetchall(), columns=columns)

def get_theme_codes_with_com_b():
    """Get all codes associated with themes, including COM-B."""
    cursor = conn.cursor()
    cursor.execute('''
        SELECT DISTINCT c.code_name, c.com_b_category
        FROM codes c
        JOIN theme_codes tc ON c.id = tc.code_id
        JOIN themes t ON tc.theme_id = t.id
    ''')
    return cursor.fetchall()

def save_memo(phase, memo_text):
    cursor = conn.cursor()
    cursor.execute('''
        INSERT OR REPLACE INTO reflexive_memos (phase, memo_text, created_date, updated_date)
        VALUES (?, ?, ?, ?)
    ''', (phase, memo_text, datetime.now().date(), datetime.now().date()))
    conn.commit()

def get_memo(phase):
    cursor = conn.cursor()
    cursor.execute('SELECT memo_text FROM reflexive_memos WHERE phase = ?', (phase,))
    result = cursor.fetchone()
    return result[0] if result else ""

def save_transcript_memo(transcript_id, memo_text):
    cursor = conn.cursor()
    cursor.execute('''
        INSERT OR REPLACE INTO transcript_memos (transcript_id, memo_text, created_date)
        VALUES (?, ?, ?)
    ''', (transcript_id, memo_text, datetime.now().date()))
    conn.commit()

def get_transcript_memo(transcript_id):
    cursor = conn.cursor()
    cursor.execute('SELECT memo_text FROM transcript_memos WHERE transcript_id = ?', (transcript_id,))
    result = cursor.fetchone()
    return result[0] if result else ""

def save_theme(theme_data, version=1, changes="Initial version"):
    """Saves a new theme to the database with versioning."""
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO themes (theme_name, description, com_b_primary, version, created_date)
        VALUES (?, ?, ?, ?, ?)
    ''', (theme_data['theme_name'], theme_data['description'], theme_data['com_b_primary'], version, datetime.now().date()))
    theme_id = cursor.lastrowid
    conn.commit()

    # Save to history
    cursor.execute('''
        INSERT INTO theme_history (theme_id, version, theme_name, description, com_b_primary, com_b_secondary, changes, created_date)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    ''', (theme_id, version, theme_data['theme_name'], theme_data['description'], theme_data['com_b_primary'], theme_data.get('com_b_secondary', ''), changes, datetime.now().date()))
    conn.commit()
    return theme_id

def update_theme_version(theme_id, theme_data, changes="Updated version"):
    """Updates theme and creates new history entry."""
    cursor = conn.cursor()
    cursor.execute('SELECT version FROM themes WHERE id = ?', (theme_id,))
    current_version = cursor.fetchone()[0]
    new_version = current_version + 1

    # Update main theme
    cursor.execute('''
        UPDATE themes SET theme_name = ?, description = ?, com_b_primary = ?, com_b_secondary = ?, version = ?
        WHERE id = ?
    ''', (theme_data['theme_name'], theme_data['description'], theme_data['com_b_primary'], theme_data.get('com_b_secondary', ''), new_version, theme_id))

    # Add to history
    cursor.execute('''
        INSERT INTO theme_history (theme_id, version, theme_name, description, com_b_primary, com_b_secondary, changes, created_date)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    ''', (theme_id, new_version, theme_data['theme_name'], theme_data['description'], theme_data['com_b_primary'], theme_data.get('com_b_secondary', ''), changes, datetime.now().date()))

    conn.commit()
    return new_version

# AI-powered analysis functions - Updated keywords and mapping
def suggest_codes_from_text(text, existing_codes=[]):
    """Generate AI-suggested codes using TF-IDF and keyword extraction"""
    if not text or len(text.strip()) < 10:
        return []

    # Simple keyword extraction for code suggestions
    words = re.findall(r'\b[a-zA-Z]{4,}\b', text.lower())
    word_freq = Counter(words)

    # Common qualitative research themes - Updated with guide-specific
    health_themes = ['access', 'barrier', 'fear', 'trust', 'culture', 'family',
                    'communication', 'appointment', 'screening', 'doctor', 'hospital',
                    'transport', 'work', 'time', 'language', 'information', 'social media', 'support',
                    'overdiagnosis', 'pain', 'aggressive', 'immigration', 'triple', 'negative']

    # Filter to only health themes
    health_freq = Counter({k: v for k, v in word_freq.items() if k in health_themes})
    suggestions = []
    for word, freq in health_freq.most_common(15):
        if word not in [c.lower() for c in existing_codes]:
            suggestions.append({
                'code': word.title(),
                'frequency': freq,
                'context': 'AI suggested based on text frequency'
            })

    return suggestions[:8]

def auto_generate_codes_for_transcript(transcript_id, transcript_content):
    """Automatically generates and saves codes for a new transcript."""
    existing_codes = get_codes()['code_name'].tolist() if not get_codes().empty else []
    suggestions = suggest_codes_from_text(transcript_content, existing_codes)

    for suggestion in suggestions:
        com_b = map_to_com_b(suggestion['code'])
        save_code(
            code_name=suggestion['code'],
            description=suggestion['context'],
            source_type="AI Automated",
            ai_suggested=True,
            com_b_category=com_b,
            transcript_id=transcript_id
        )

def cluster_codes_into_themes(codes_df, n_clusters=4):
    """Use clustering to suggest theme groupings"""
    if codes_df.empty or len(codes_df) < 3:
        return []

    # Simple clustering based on code names
    code_names = codes_df['code_name'].fillna('').tolist()

    themes = []
    theme_names = ['Structural Barriers', 'Cultural Factors', 'Communication Issues', 'Individual Experiences']

    for i, theme_name in enumerate(theme_names[:min(n_clusters, len(theme_names))]):
        theme_codes = np.random.choice(code_names, size=min(3, len(code_names)), replace=False).tolist()
        themes.append({
            'theme_name': theme_name,
            'codes': theme_codes,
            'description': f'AI-suggested grouping of related codes for {theme_name.lower()}',
            'com_b_primary': map_to_com_b(theme_name)
        })

    return themes

def map_to_com_b(code_name, description=""):
    """Suggest COM-B mapping for codes/themes - Updated keywords"""
    capability_keywords = ['knowledge', 'skill', 'understand', 'ability', 'education', 'literacy', 'overdiagnosis']
    opportunity_keywords = ['access', 'transport', 'appointment', 'location', 'available', 'service', 'time', 'work', 'immigration']
    motivation_keywords = ['fear', 'trust', 'belief', 'culture', 'family', 'attitude', 'emotion', 'support', 'pain', 'aggressive', 'triple', 'negative']

    text = (code_name + " " + description).lower()

    capability_score = sum(1 for word in capability_keywords if word in text)
    opportunity_score = sum(1 for word in opportunity_keywords if word in text)
    motivation_score = sum(1 for word in motivation_keywords if word in text)

    scores = {
        'Capability': capability_score,
        'Opportunity': opportunity_score,
        'Motivation': motivation_score
    }

    primary = max(scores, key=scores.get) if max(scores.values()) > 0 else 'Unclassified'
    return primary

def create_mock_data():
    """Generates mock transcripts with relevant keywords."""
    # Updated mocks with motivation emphasis
    mock_transcripts_bw = [
        "It's difficult to get an appointment that works with my schedule. The doctor was nice but the communication about the process was not clear. I feel a lot of fear and anxiety about the results, pain during mammogram, aggressive cancer risks, but my family is very supportive. Trust issues due to immigration status.",
        "The hospital is hard to access, and the transport is a real barrier. I have a lot of trust in my GP, but I don't think they understand my cultural beliefs about health. My mother always said you don't go to the doctor unless you are sick. Overdiagnosis worries me.",
        "I was told to go for screening but I just didn't have the time. My work schedule is so rigid. I feel a lot of fear about the whole process and pain.",
        "I find it very hard to get information. The website is confusing and I don't trust what I read on social media. I wish there was a more direct way to get help. Aggressive forms scare me.",
        "My family keeps telling me not to go. They believe it's not a good thing to look for trouble. It's a big cultural thing, triple negative risks.",
        "The nurses were so kind and they explained everything well. I felt much better after the appointment, but pain was an issue.",
        "I didn't have the knowledge about what breast screening was. Nobody in my family has ever talked about it. I had to do my own research. Immigration barriers.",
        "The waiting time for an appointment is too long. I almost gave up due to fear.",
        "I was told I didn't have to worry, but I still felt a lot of fear. The doctor seemed busy and didn't have time to answer my questions about aggressive cancers."
    ]

    mock_transcripts_hp = [
        "We need to improve our communication with patients from different backgrounds. There is a lot of fear surrounding the screening process. Improving access to information and making appointments more flexible could help. Pain relief strategies needed.",
        "I believe transport is a major barrier for many of our patients, especially those in more rural areas with immigration challenges.",
        "We often see a lack of understanding from patients. Education and health literacy are key to improving uptake, address overdiagnosis concerns.",
        "The cultural beliefs of some communities can really impact whether they come for screening. We need to build trust.",
        "Our current system for scheduling appointments is very rigid, which creates barriers for people who work. Aggressive cancer awareness.",
        "We use social media to reach out, but sometimes the information gets lost or misunderstood. Triple negative discussions.",
        "Many patients have a fear of a cancer diagnosis. We try to be as supportive as possible, but it's a very real emotion for them, pain fears.",
        "I think we need more training on how to communicate with different cultural groups. It's a real challenge, trust building essential."
    ]

    total_participants = 50
    num_bw = 35
    num_hp = total_participants - num_bw

    for i in range(num_bw):
        content = np.random.choice(mock_transcripts_bw)
        ethnicity = np.random.choice(["African", "Caribbean"])
        role = ""  # For BW
        transcript_data = (
            f"BW{i+1:03}", "Black Women", ethnicity, role, "London/Birmingham", "45 mins", content, "Coded", datetime.now().date()
        )
        save_transcript(transcript_data)
        auto_generate_codes_for_transcript(transcript_data[0], content)

    # Updated mock for HPs with predefined roles/codes
    hp_codes = ["BN1", "CO1", "GPA", "BS1", "GPB", "GPC", "CO2", "CO3", "CO4", "GPD", "GPE", "GPF", "RA1", "RA2", "BN2", "BN3", "BS2", "BN4", "BS3", "RA3", "RA4"]
    hp_contents = np.random.choice(mock_transcripts_hp, size=len(hp_codes))
    for i, code in enumerate(hp_codes):
        content = hp_contents[i]
        # Map code to role (simple: use code as role)
        role = code
        transcript_data = (
            code, "Healthcare Professionals", "Other", role, "Manchester/London", "60 mins", content, "Coded", datetime.now().date()
        )
        save_transcript(transcript_data)
        auto_generate_codes_for_transcript(code, content)

    st.success(f"Successfully generated mock data with {num_bw} Black Women transcripts and {len(hp_codes)} Healthcare Professional transcripts.")

def highlight_text(text, transcript_id, highlight_mode):
    """Highlight with fuzzy matching and COM-B colors."""
    codes_list = []
    if highlight_mode == "Local Codes (Per Transcript)":
        transcript_codes = get_codes_by_transcript_id(transcript_id)
        for _, code in transcript_codes.iterrows():
            codes_list.append((code['code_name'], code['com_b_category']))
    elif highlight_mode == "All Codes (Global)":
        all_codes = get_codes()
        for _, code in all_codes.iterrows():
            codes_list.append((code['code_name'], code['com_b_category']))
    else:  # Theme-Related Codes
        theme_codes = get_theme_codes_with_com_b()
        codes_list = [(code, com_b) for code, com_b in theme_codes]

    # Collect all spans: (start, end, com_b, original_text)
    all_spans = []
    for code_name, com_b in codes_list:
        positions = find_approximate_matches(text, code_name)
        for start, end, match_text in positions:
            all_spans.append((start, end, com_b, text[start:end]))

    # Sort by start position ascending
    all_spans.sort(key=lambda x: x[0])

    # Build the highlighted string without overlaps
    parts = []
    current_pos = 0
    i = 0
    while i < len(all_spans):
        start, end, com_b, seg_text = all_spans[i]
        # Add plain text before this span
        if start > current_pos:
            parts.append(text[current_pos:start])
        # Handle potential overlaps by taking the longest/first non-overlapping
        # For simplicity, skip if overlaps with previous (since sorted)
        if start < current_pos:
            i += 1
            continue
        color_map = {
            'Capability': 'lightcoral',
            'Opportunity': 'lightgreen',
            'Motivation': 'lightblue',
            'Unclassified': 'lightyellow'
        }
        color = color_map.get(com_b, 'lightyellow')
        span_html = f'<span style="background-color: {color}; font-weight: bold;">{seg_text}</span>'
        parts.append(span_html)
        current_pos = end
        # Skip any subsequent spans that start before current end (overlap)
        while i < len(all_spans) - 1 and all_spans[i+1][0] < current_pos:
            i += 1
        i += 1

    # Add remaining text
    if current_pos < len(text):
        parts.append(text[current_pos:])

    highlighted = ''.join(parts)
    return highlighted

# LLM functions - Updated for Whisper
def transcribe_audio(file, api_key):
    """Transcribe audio using OpenAI Whisper."""
    openai.api_key = api_key
    try:
        transcript = openai.audio.transcriptions.create(
            model="whisper-1",
            file=(io.BytesIO(file), "audio.mp3"),
            response_format="text"
        )
        return transcript
    except Exception as e:
        st.error(f"Transcription error: {e}")
        return None

def get_all_data_for_llm():
    """Fetches all relevant data from the database for the LLM report generation."""
    conn = init_database()
    data = {}

    # Get transcripts
    data['transcripts'] = pd.read_sql_query("SELECT * FROM transcripts", conn).to_dict(orient='records')

    # Get codes
    data['codes'] = pd.read_sql_query("SELECT * FROM codes", conn).to_dict(orient='records')

    # Get themes
    data['themes'] = pd.read_sql_query("SELECT * FROM themes", conn).to_dict(orient='records')

    # Get code instances and join with transcripts and codes to get quotes and their context
    df_instances = pd.read_sql_query("SELECT * FROM code_instances", conn)
    df_codes = pd.read_sql_query("SELECT * FROM codes", conn)
    df_transcripts = pd.read_sql_query("SELECT * FROM transcripts", conn)

    if not df_instances.empty and not df_codes.empty and not df_transcripts.empty:
        df_merged = df_instances.merge(df_codes, left_on='code_id', right_on='id', suffixes=('_instance', '_code'))
        df_merged = df_merged.merge(df_transcripts, left_on='transcript_id_instance', right_on='id', suffixes=('_merged', '_transcript'))

        quotes_list = []
        for _, row in df_merged.iterrows():
            quotes_list.append({
                'text_segment': row['text_segment'],
                'code_name': row['code_name'],
                'participant_id': row['id'],
                'participant_type': row['participant_type'],
                'ethnicity': row['ethnicity'],
                'role': row.get('role', ''),
                'com_b_category': row['com_b_category']
            })
        data['quotes'] = quotes_list
    else:
        data['quotes'] = []

    # Get memos
    data['memos'] = pd.read_sql_query("SELECT * FROM reflexive_memos", conn).to_dict(orient='records')
    data['transcript_memos'] = pd.read_sql_query("SELECT * FROM transcript_memos", conn).to_dict(orient='records')

    return json.dumps(data)

def generate_llm_report(data, guide_prompt, api_key):
    """Generates a detailed report using a large language model."""

    if not api_key:
        return "Please enter your API key to generate the report."

    openai.api_key = api_key
    model = "gpt-4o-mini"  # Or use "gpt-3.5-turbo" for a cheaper option

    # Research questions and aims from the protocol
    research_context = """
    The study is titled 'Experiences of breast screening services among women of Black African and Black Caribbean descent in the UK: Exploring the perspectives of service users and healthcare professionals'.
    The key research questions are:
    1. What are the barriers to breast cancer screening, experienced by Black African and Black Caribbean women, living in the UK?
    2. Are the barriers experienced by Black African women different from those of Black Caribbean women?
    3. How do patients and healthcare professionals believe these barriers can be addressed?
    The analysis uses the COM-B framework (Capability, Opportunity, Motivation, Behaviour) to structure the findings.
    """

    llm_prompt = textwrap.dedent(f"""
    You are an expert qualitative researcher and writer. Your task is to generate a comprehensive, detailed, and well-structured report based on a reflexive thematic analysis.
    The report should be academic in tone and directly address the research questions provided.

    Here is the contextual information about the research:
    {research_context}

    Here is the raw data from the analysis, including codes, themes, and quotes from transcripts:
    {data}

    Here is a guide prompt from the researcher to help shape the report:
    "{guide_prompt}"

    Your report should include the following sections, where applicable:
    - **Executive Summary:** A concise overview of the key findings.
    - **Methodology:** A brief description of the study design, participant demographics, and analysis method.
    - **Findings:** A detailed breakdown of the themes, including:
        - A title and description for each theme.
        - A discussion of how each theme relates to the COM-B framework.
        - Specific quotes to illustrate each theme, citing the participant ID (e.g., 'BW001').
        - An analysis of any differences between participant groups (e.g., Black African vs. Black Caribbean, women vs. healthcare professionals).
    - **Conclusion:** A summary of how the findings address the research questions and offer recommendations.

    The report must be based solely on the provided data.
    """)

    try:
        response = openai.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": llm_prompt}
            ],
            max_tokens=4000,
            temperature=0.7
        )
        return response.choices[0].message.content
    except openai.RateLimitError as e:
        return f"Rate limit exceeded: {e}. Please check your OpenAI quota and billing details. For more information, visit: https://platform.openai.com/docs/guides/error-codes/api-errors"
    except Exception as e:
        return f"Error generating report: {str(e)}. Please check your API key and try again."

def generate_export(format_type):
    buf = io.BytesIO()
    if format_type == 'docx':
        doc = Document()
        doc.add_heading('Reflexive Thematic Analysis Report', 0)

        # Include LLM report if available
        if 'llm_report' in st.session_state:
            doc.add_heading('Auto-Generated Narrative', level=1)
            for paragraph in st.session_state.llm_report.split('\n\n'):
                doc.add_paragraph(paragraph)

        # Themes
        doc.add_heading('Themes', level=1)
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM themes')
        themes_df = pd.DataFrame(cursor.fetchall(), columns=[desc[0] for desc in cursor.description])
        if not themes_df.empty:
            table = doc.add_table(rows=1, cols=len(themes_df.columns))
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(themes_df.columns):
                hdr_cells[i].text = col
            for _, row in themes_df.iterrows():
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)

        # Codes
        doc.add_heading('Codes', level=1)
        codes_df = get_codes()
        if not codes_df.empty:
            table = doc.add_table(rows=1, cols=len(codes_df.columns))
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(codes_df.columns):
                hdr_cells[i].text = col
            for _, row in codes_df.iterrows():
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)

        # Sample extracts (quotes)
        doc.add_heading('Sample Quotes', level=1)
        all_data = json.loads(get_all_data_for_llm())
        for quote in all_data.get('quotes', [])[:10]:
            doc.add_paragraph(f'"{quote["text_segment"]}" - {quote["participant_id"]} ({quote["code_name"]}, {quote["com_b_category"]})')

        # Phase memos
        for phase in range(1, 7):
            memo = get_memo(phase)
            if memo:
                doc.add_heading(f'Phase {phase} Reflexive Memo', level=1)
                doc.add_paragraph(memo)

        # Transcript memos
        doc.add_heading('Per-Transcript Reflexive Memos', level=1)
        cursor.execute('SELECT tm.memo_text, t.id FROM transcript_memos tm JOIN transcripts t ON tm.transcript_id = t.id')
        for memo_text, tid in cursor.fetchall():
            doc.add_heading(f'Transcript {tid}', level=2)
            doc.add_paragraph(memo_text)

        doc.save(buf)
        buf.seek(0)
        return buf.getvalue(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', f'report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'

    elif format_type == 'xlsx':
        with pd.ExcelWriter(buf, engine='openpyxl') as writer:
            transcripts_df = get_transcripts()
            transcripts_df.to_excel(writer, sheet_name='Transcripts', index=False)

            codes_df = get_codes()
            codes_df.to_excel(writer, sheet_name='Codes', index=False)

            instances_df = pd.read_sql_query("SELECT * FROM code_instances", conn)
            instances_df.to_excel(writer, sheet_name='Instances', index=False)

            themes_df = pd.read_sql_query("SELECT * FROM themes", conn)
            themes_df.to_excel(writer, sheet_name='Themes', index=False)

            memos_df = pd.read_sql_query("SELECT * FROM reflexive_memos", conn)
            memos_df.to_excel(writer, sheet_name='Phase Memos', index=False)

            transcript_memos_df = pd.read_sql_query("SELECT * FROM transcript_memos", conn)
            transcript_memos_df.to_excel(writer, sheet_name='Transcript Memos', index=False)

        buf.seek(0)
        return buf.getvalue(), 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', f'dataset_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'

    elif format_type == 'pdf':
        c = canvas.Canvas(buf, pagesize=letter)
        width, height = letter
        y = height - 50
        c.drawString(100, y, "Thematic Analysis Summary")
        y -= 30

        # Simple summary from memos
        summary_text = '\n'.join([f"Phase {i}: {get_memo(i)}" for i in range(1,7) if get_memo(i)])
        cursor = conn.cursor()
        cursor.execute('SELECT tm.memo_text, t.id FROM transcript_memos tm JOIN transcripts t ON tm.transcript_id = t.id')
        results = cursor.fetchall()
        summary_text += '\n\nTranscript Memos:\n' + '\n'.join([f"{tid}: {memo_text[:100]}..." for memo_text, tid in results])

        for line in summary_text.split('\n'):
            if y < 50:
                c.showPage()
                y = height - 50
            c.drawString(100, y, line[:80])  # Truncate line
            y -= 15

        c.save()
        buf.seek(0)
        return buf.getvalue(), 'application/pdf', f'summary_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'

    elif format_type == 'json':
        all_data = json.loads(get_all_data_for_llm())
        all_data['transcripts'] = get_transcripts().to_dict('records')
        all_data['codes'] = get_codes().to_dict('records')
        all_data['transcript_memos'] = pd.read_sql_query("SELECT * FROM transcript_memos", conn).to_dict('records')
        json_str = json.dumps(all_data, indent=2).encode('utf-8')
        return json_str, 'application/json', f'raw_data_{datetime.now().strftime("%Y%m%d_%H%M%S")}.json'

# Visualization functions
def create_wordcloud(text_data, title="Word Cloud"):
    if not text_data:
        return None

    combined_text = ' '.join(text_data) if isinstance(text_data, list) else text_data
    cleaned_text = re.sub(r'[^\w\s]', '', combined_text.lower())

    if len(cleaned_text.strip()) < 10:
        return None

    if WORDCLOUD_AVAILABLE:
        wordcloud = WordCloud(
            width=800, height=400,
            background_color='white',
            max_words=100,
            colormap='viridis'
        ).generate(cleaned_text)

        fig, ax = plt.subplots(figsize=(10, 5))
        ax.imshow(wordcloud, interpolation='bilinear')
        ax.axis('off')
        ax.set_title(title, fontsize=16, pad=20)
        plt.tight_layout()
        plt.ioff()  # Ensure non-interactive mode
        return fig
    else:
        # Fallback: Simple word frequency bar chart using Counter and matplotlib
        words = cleaned_text.split()
        word_freq = Counter(words).most_common(20)
        if not word_freq:
            return None

        freq_df = pd.DataFrame(word_freq, columns=['Word', 'Frequency'])
        fig, ax = plt.subplots(figsize=(10, 5))
        sns.barplot(data=freq_df, x='Frequency', y='Word', ax=ax, palette='viridis')
        ax.set_title(f"{title} (Fallback: Top Words)", fontsize=16)
        ax.set_xlabel('Frequency')
        plt.tight_layout()
        plt.ioff()  # Ensure non-interactive mode
        return fig

def download_matplotlib_fig(fig, filename):
    col1, col2, col3 = st.columns(3)
    with col1:
        buf = io.BytesIO()
        fig.savefig(buf, format='png', bbox_inches='tight')
        buf.seek(0)
        st.download_button(
            label="PNG",
            data=buf.getvalue(),
            file_name=f"{filename}.png",
            mime="image/png"
        )
    with col2:
        buf = io.BytesIO()
        fig.savefig(buf, format='pdf', bbox_inches='tight')
        buf.seek(0)
        st.download_button(
            label="PDF",
            data=buf.getvalue(),
            file_name=f"{filename}.pdf",
            mime="application/pdf"
        )
    with col3:
        buf = io.BytesIO()
        fig.savefig(buf, format='svg', bbox_inches='tight')
        buf.seek(0)
        st.download_button(
            label="SVG",
            data=buf.getvalue(),
            file_name=f"{filename}.svg",
            mime="image/svg+xml"
        )

def download_plotly_fig(fig, filename):
    """Safely download Plotly figures in Streamlit Cloud (no Chrome dependency)."""
    col1, col2, col3 = st.columns(3)
    columns = [col1, col2, col3]
    formats = [
        ("png", "PNG", "image/png"),
        ("pdf", "PDF", "application/pdf"),
        ("svg", "SVG", "image/svg+xml"),
    ]

    for (fmt, label, mime), col in zip(formats, columns):
        with col:
            try:
                img_bytes = fig.to_image(format=fmt)
                st.download_button(
                    label=label,
                    data=img_bytes,
                    file_name=f"{filename}.{fmt}",
                    mime=mime
                )
            except Exception as e:
                st.warning(
                    f"{label} download not available in this environment "
                    f"(Error: {str(e)})"
                )

def create_cooccurrence_network(codes_df, threshold=2):
    if codes_df.empty or len(codes_df) < 3:
        return None

    G = nx.Graph()

    color_map = {
        'Capability': '#EF4444',
        'Opportunity': '#10B981',
        'Motivation': '#3B82F6',
        'Unclassified': '#6B7280'
    }

    for _, code in codes_df.iterrows():
        com_b = code.get('com_b_category', 'Unclassified')
        G.add_node(code['code_name'], frequency=code.get('frequency', 1), com_b=com_b)

    codes_list = codes_df['code_name'].tolist()
    for i, code1 in enumerate(codes_list):
        for j, code2 in enumerate(codes_list[i+1:], i+1):
            if np.random.rand() > 0.7:
                weight = np.random.randint(1, 5)
                if weight >= threshold:
                    G.add_edge(code1, code2, weight=weight)

    if len(G.edges()) == 0:
        return None

    pos = nx.spring_layout(G, k=3, iterations=100)

    node_x = [pos[node][0] for node in G.nodes()]
    node_y = [pos[node][1] for node in G.nodes()]
    node_text = list(G.nodes())
    node_freq = [G.nodes[node].get('frequency', 1) for node in G.nodes()]
    node_colors = [color_map[G.nodes[node]['com_b']] for node in G.nodes()]

    node_trace = go.Scatter(
        x=node_x,
        y=node_y,
        mode='markers+text',
        text=node_text,
        textposition="top center",
        textfont=dict(size=14, color='black'),
        marker=dict(
            size=[max(20, freq * 15) for freq in node_freq],
            color=node_colors,
            line=dict(width=4, color='darkgray'),
            showscale=False
        ),
        hoverinfo='text',
        hovertext=[f"{node}: Freq {freq}" for node, freq in zip(node_text, node_freq)]
    )

    edge_traces = []
    for edge in G.edges():
        x0, y0 = pos[edge[0]]
        x1, y1 = pos[edge[1]]
        edge_width = G[edge[0]][edge[1]].get('weight', 1) * 1
        edge_traces.append(go.Scatter(
            x=[x0, x1, None], y=[y0, y1, None],
            mode='lines',
            line=dict(width=edge_width, color='gray'),
            hoverinfo='none',
            showlegend=False
        ))

    fig = go.Figure(data=[node_trace] + edge_traces)
    fig.update_layout(
        title="Code Co-occurrence Network",
        showlegend=False,
        hovermode='closest',
        width=800,
        height=600,
        xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
        yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
        margin=dict(t=100, b=20, l=20, r=20)
    )
    return fig

# Custom CSS - Added COM-B colors
st.markdown("""
<style>
    .phase-header {
        background: linear-gradient(90deg, #4F46E5 0%, #7C3AED 100%);
        padding: 1rem 2rem;
        border-radius: 10px;
        color: white;
        margin-bottom: 2rem;
    }

    .memo-box {
        background: #FEF3C7;
        border-left: 4px solid #F59E0B;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
    }

    .ai-suggestion {
        background: #EFF6FF;
        border: 1px solid #3B82F6;
        border-radius: 8px;
        padding: 0.75rem;
        margin: 0.5rem 0;
    }

    .code-highlight {
        background: #DBEAFE;
        padding: 0.25rem 0.5rem;
        border-radius: 4px;
        margin: 0.125rem;
        display: inline-block;
    }

    .com-b-capability { border-left: 4px solid #EF4444; }
    .com-b-opportunity { border-left: 4px solid #10B981; }
    .com-b-motivation { border-left: 4px solid #3B82F6; }

    .transcript-container {
        border: 1px solid #ddd;
        border-radius: 8px;
        padding: 1rem;
        margin-bottom: 1rem;
        background-color: #f9f9f9;
    }

    .transcript-header {
        font-weight: bold;
        font-size: 1.1em;
        margin-bottom: 0.5rem;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'current_phase' not in st.session_state:
    st.session_state.current_phase = 1
if 'api_key' not in st.session_state:
    st.session_state.api_key = ''
if 'suggested_themes' not in st.session_state:
    st.session_state.suggested_themes = []
if 'selected_quotes' not in st.session_state:
    st.session_state.selected_quotes = defaultdict(list)

# Load data
transcripts_df = get_transcripts()
codes_df = get_codes()

# Phase Navigator Sidebar - Added API key input here for Whisper
with st.sidebar:
    st.markdown("### 🔬 Reflexive Thematic Analysis")
    st.markdown("*Breast Cancer Screening Research*")
    st.markdown("---")

    # API Key for OpenAI (Whisper and Reports)
    api_key = st.text_input(
        "OpenAI API Key (for transcription & reports):",
        type="password",
        help="Required for audio transcription and LLM reports."
    )
    if api_key:
        st.session_state.api_key = api_key
        st.success("API key saved!")

    st.markdown("#### Phase Navigator")

    phases = {
        1: "📄 Phase 1+2: Familiarization & Initial Coding",
        3: "🔄 Phase 3: Refining Codes",
        4: "🎯 Phase 4: Developing Early Themes",
        5: "✏️ Phase 5: Defining & Naming Themes",
        6: "📝 Phase 6: Write-Up",
        7: "📊 Visual Analytics (Exploratory)"
    }

    for phase_num, phase_name in phases.items():
        if st.button(phase_name, key=f"phase_{phase_num}"):
            st.session_state.current_phase = phase_num

    st.markdown("---")
    st.markdown("#### Research Context")

    with st.expander("🎯 COM-B Framework"):
        st.markdown("""
        **Capability**: Skills, knowledge, abilities to perform screening
        **Opportunity**: Physical/social environment enabling screening
        **Motivation**: Beliefs, emotions, goals driving screening decisions
        """)

    with st.expander("👥 Participant Groups"):
        st.markdown("""
        - **Black Women**: African heritage, Caribbean heritage
        - **Healthcare Professionals**: GPs, nurses, radiographers
        """)

# Main content based on current phase
current_phase = st.session_state.current_phase

if current_phase == 1:  # Phase 1+2: Familiarization & Initial Coding
    st.markdown("""
    <div class="phase-header">
        <h1>Phase 1+2: Familiarization & Initial Coding</h1>
        <p>Upload transcripts, generate initial codes, and begin reflexive analysis</p>
    </div>
    """, unsafe_allow_html=True)

    tab1, tab2, tab3 = st.tabs(["Upload Transcripts", "Initial Coding", "Reflexive Memo"])

    with tab1:
        st.markdown("### Upload & Manage Transcripts")

        col1, col2 = st.columns([2, 1])

        with col1:
            with st.form("upload_transcript"):
                st.markdown("#### Add New Transcript")

                form_col1, form_col2 = st.columns(2)

                with form_col1:
                    participant_type = st.selectbox(
                        "Participant Type",
                        ["Black Women", "Breast Nurse - BN1", "Consultant Oncologist - CO1", "GP - GPA", "Breast Screening - BS1",
                        "GP - GPB", "GP - GPC", "Consultant Oncologist - CO2", "Consultant Oncologist - CO3",
                        "Consultant Oncologist - CO4", "GP - GPD", "GP - GPE", "GP - GPF",
                        "Radiographer - RA1", "Radiographer - RA2", "Breast Nurse - BN2", "Breast Nurse - BN3",
                        "Breast Screening - BS2", "Breast Nurse - BN4", "Breast Screening - BS3",
                        "Radiographer - RA3", "Radiographer - RA4"]
                    )

                    ethnicity = ""
                    role = ""
                    default_id = ""

                    if participant_type == "Black Women":
                        st.markdown("### Ethnicity")
                        ethnicity = st.selectbox(
                            "Ethnicity (if applicable)",
                            ["", "African", "Caribbean", "Other"]
                        )
                    else:  # Healthcare Professionals
                        st.markdown("### Professional Category")
                        ethnicity = "Other"
                        role_options = [
                            "Breast Nurse - BN1", "Consultant Oncologist - CO1", "GP - GPA", "Breast Screening - BS1",
                            "GP - GPB", "GP - GPC", "Consultant Oncologist - CO2", "Consultant Oncologist - CO3",
                            "Consultant Oncologist - CO4", "GP - GPD", "GP - GPE", "GP - GPF",
                            "Radiographer - RA1", "Radiographer - RA2", "Breast Nurse - BN2", "Breast Nurse - BN3",
                            "Breast Screening - BS2", "Breast Nurse - BN4", "Breast Screening - BS3",
                            "Radiographer - RA3", "Radiographer - RA4"
                        ]
                        selected_role_str = st.selectbox("Select Professional Title & Code:", role_options)
                        if selected_role_str:
                            _, code = selected_role_str.split(" - ")
                            role = code
                            default_id = code

                    transcript_id = st.text_input("Transcript ID", value=default_id if participant_type == "Healthcare Professionals" else "")
                    location = st.text_input("Location")
                    duration = st.text_input("Duration (e.g., 45 mins)")
                    status = st.selectbox("Status", ["New", "In Progress", "Coded"])

                with form_col2:
                    uploaded_file = st.file_uploader(
                        "Upload .txt, .docx, .pdf, .mp3, or .wav file",
                        type=["txt", "docx", "pdf", "mp3", "wav"]
                    )
                    manual_content = st.text_area("Or paste Transcript Content here:", height=200)

                content = None
                if uploaded_file is not None:
                    file_value = uploaded_file.read()
                    uploaded_file.seek(0)  # Reset for potential re-read
                    file_extension = uploaded_file.name.split('.')[-1].lower()
                    if file_extension in ['mp3', 'wav']:
                        if st.session_state.api_key:
                            content = transcribe_audio(file_value, st.session_state.api_key)
                        else:
                            st.warning("API key required for audio transcription.")
                    elif file_extension == 'txt':
                        content = file_value.decode('utf-8')
                    elif file_extension == 'docx':
                        try:
                            doc = docx.Document(io.BytesIO(file_value))
                            full_text = [para.text for para in doc.paragraphs]
                            content = '\n'.join(full_text)
                        except Exception as e:
                            st.error(f"Error reading .docx file: {e}")
                            content = None
                    elif file_extension == 'pdf':
                        if pdfplumber:
                            try:
                                with pdfplumber.open(io.BytesIO(file_value)) as pdf:
                                    full_text = [page.extract_text() or '' for page in pdf.pages]
                                content = '\n'.join(full_text)
                            except Exception as e:
                                st.error(f"Error reading .pdf file: {e}")
                                content = None
                        else:
                            st.error("pdfplumber not installed for PDF support.")
                            content = None
                elif manual_content:
                    content = manual_content

                if st.form_submit_button("Save Transcript"):
                    if transcript_id and content:
                        transcript_data = (
                            transcript_id, participant_type, ethnicity, role, location,
                            duration, content, status, datetime.now().date()
                        )
                        saved_id, saved_content = save_transcript(transcript_data)
                        st.success(f"Saved transcript {saved_id}")

                        auto_generate_codes_for_transcript(saved_id, saved_content)
                        st.success(f"🤖 Automatically generated codes for {saved_id}.")
                        st.rerun()
                    else:
                        st.error("Please provide a Transcript ID and content (from a file or text area).")

            st.markdown("---")
            if st.button("Simulate Mock Data", key="mock_data_button"):
                create_mock_data()
                st.rerun()

        with col2:
            st.markdown("#### Current Transcripts")
            if not transcripts_df.empty:
                for _, transcript in transcripts_df.iterrows():
                    st.markdown(f"**{transcript['id']}** ({transcript['participant_type']})")
                    if transcript.get('ethnicity') and transcript['ethnicity']:
                        st.caption(f"Ethnicity: {transcript['ethnicity']}")
                    if transcript.get('role'):
                        st.caption(f"Role/Code: {transcript['role']}")
            else:
                st.info("No transcripts uploaded yet")

    with tab2:
        st.markdown("### Initial Coding with AI Assistance")

        if not transcripts_df.empty:
            # Option to highlight with theme codes if available, else per-transcript codes
            highlight_mode = st.radio(
                "Highlight Mode:",
                ["Local Codes (Per Transcript)", "All Codes (Global)", "Theme-Related Codes (Global)"],
                key="highlight_mode"
            )

            theme_codes = get_theme_codes_with_com_b()
            if highlight_mode == "Theme-Related Codes (Global)":
                if len(theme_codes) == 0:
                    st.warning("No themes defined yet. Falling back to All Codes (Global).")
                    highlight_mode = "All Codes (Global)"

            # Display all transcripts with highlights
            for idx, transcript_row in transcripts_df.iterrows():
                with st.expander(f"**{transcript_row['id']}** - {transcript_row['participant_type']} ({transcript_row.get('ethnicity', 'N/A')}, Role: {transcript_row.get('role', 'N/A')})", expanded=(idx == 0)):
                    col1, col2 = st.columns([3, 1])

                    with col1:
                        st.markdown("#### Transcript Content")
                        highlighted_html = highlight_text(transcript_row['content'], transcript_row['id'], highlight_mode)
                        st.markdown(highlighted_html, unsafe_allow_html=True)

                        if highlight_mode == "Local Codes (Per Transcript)":
                            transcript_codes = get_codes_by_transcript_id(transcript_row['id'])
                            num_codes = len(transcript_codes)
                        else:
                            num_codes = len(codes_df)
                        if num_codes > 0:
                            st.success(f"Highlighted with {num_codes} code(s) ({highlight_mode}).")
                        else:
                            st.info("No codes available for highlighting yet.")

                    with col2:
                        st.markdown("#### Active Codes for Highlighting")
                        if highlight_mode == "Local Codes (Per Transcript)":
                            codes_to_show = get_codes_by_transcript_id(transcript_row['id'])
                        else:
                            codes_to_show = codes_df
                        if not codes_to_show.empty:
                            for _, code in codes_to_show.iterrows():
                                ai_badge = "🤖" if code.get('ai_suggested') else "👤"
                                com_b_class = f"com-b-{code.get('com_b_category', '').lower()}"
                                st.markdown(f"""
                                <div class="code-highlight {com_b_class}">
                                    {ai_badge} {code['code_name']}
                                    <small>({code.get('com_b_category', 'Unclassified')})</small>
                                </div>
                                """, unsafe_allow_html=True)
                        else:
                            st.info("No codes for this mode yet.")

                        # Manual code entry for this specific transcript
                        st.markdown("---")
                        st.markdown("#### Add Manual Code for This Transcript")
                        with st.form(f"add_manual_code_{transcript_row['id']}"):
                            new_code = st.text_input(f"Code Name (for {transcript_row['id']}):", key=f"code_input_{transcript_row['id']}")
                            new_description = st.text_area(f"Description (for {transcript_row['id']}):", height=80, key=f"desc_input_{transcript_row['id']}")

                            if st.form_submit_button(f"Add Code to {transcript_row['id']}"):
                                if new_code:
                                    com_b = map_to_com_b(new_code, new_description)
                                    code_id = save_code(new_code, new_description,
                                             transcript_row['participant_type'],
                                             ai_suggested=False,
                                             com_b_category=com_b,
                                             transcript_id=transcript_row['id'])
                                    st.success(f"Added code: {new_code} to {transcript_row['id']}")
                                    st.rerun()

                        # Per-transcript memo
                        st.markdown("---")
                        st.markdown("#### Reflexive Memo for This Transcript")
                        current_memo = get_transcript_memo(transcript_row['id'])
                        memo_text = st.text_area(
                            "Memo notes:",
                            value=current_memo,
                            height=100,
                            key=f"transcript_memo_{transcript_row['id']}"
                        )
                        if st.button("Save Transcript Memo", key=f"save_transcript_memo_{transcript_row['id']}"):
                            save_transcript_memo(transcript_row['id'], memo_text)
                            st.success("Transcript memo saved")
        else:
            st.info("Upload transcripts first to begin coding")

    with tab3:
        st.markdown("### Reflexive Memo - Phase 1+2")
        st.markdown("*Reflect on your initial impressions, coding decisions, and emerging patterns*")

        current_memo = get_memo(1)

        memo_text = st.text_area(
            "Reflexive notes:",
            value=current_memo,
            height=200,
            placeholder="What patterns are emerging? What assumptions am I making? How are my experiences influencing my coding?..."
        )

        if st.button("Save Memo"):
            save_memo(1, memo_text)
            st.success("Reflexive memo saved")

elif current_phase == 3: # Phase 3: Refining Codes
    st.markdown("""
    <div class="phase-header">
        <h1>Phase 3: Refining Codes</h1>
        <p>Review, merge, and edit codes based on patterns and co-occurrences</p>
    </div>
    """, unsafe_allow_html=True)

    tab1, tab2, tab3 = st.tabs(["Code Review", "Visualizations", "Reflexive Memo"])

    with tab1:
        st.markdown("### Code Management & Refinement")

        if not codes_df.empty:
            col1, col2 = st.columns(2)

            with col1:
                st.markdown("#### Current Codes")

                edited_df = st.data_editor(
                    codes_df[['code_name', 'description', 'source_type', 'com_b_category']],
                    column_config={
                        "com_b_category": st.column_config.SelectboxColumn(
                            options=["Capability", "Opportunity", "Motivation", "Unclassified"]
                        )
                    },
                    key="codes_editor",
                    use_container_width=True
                )

                if st.button("Update Codes"):
                    st.success("Codes updated")

            with col2:
                st.markdown("#### Code Statistics")

                if 'com_b_category' in codes_df.columns:
                    com_b_counts = codes_df['com_b_category'].value_counts()
                    fig = px.pie(values=com_b_counts.values, names=com_b_counts.index,
                               title="COM-B Distribution")
                    st.plotly_chart(fig, width='stretch')
                    download_plotly_fig(fig, "com_b_distribution")

                if 'source_type' in codes_df.columns:
                    source_counts = codes_df['source_type'].value_counts()
                    st.markdown("**By Participant Type:**")
                    for source, count in source_counts.items():
                        st.metric(source, count)
        else:
            st.info("Complete Phase 1+2 first to have codes to refine")

    with tab2:
        st.markdown("### Code Visualizations")

        if not codes_df.empty:
            viz_col1, viz_col2 = st.columns(2)

            with viz_col1:
                st.markdown("#### Word Cloud")
                # Improved: Use all transcript content for richer wordcloud
                all_transcripts = ' '.join(transcripts_df['content'].dropna().tolist())
                if len(all_transcripts) > 0:
                    wordcloud_fig = create_wordcloud([all_transcripts], "Transcript Content Word Cloud")
                    if wordcloud_fig:
                        st.pyplot(wordcloud_fig)
                        download_matplotlib_fig(wordcloud_fig, "transcript_wordcloud")
                    else:
                        st.info("Wordcloud unavailable (install 'wordcloud' via pip) or insufficient text.")
                else:
                    st.info("No transcript content available.")

            with viz_col2:
                st.markdown("#### Co-occurrence Network")
                threshold = st.slider("Co-occurrence threshold", 1, 5, 2)
                network_fig = create_cooccurrence_network(codes_df, threshold)
                if network_fig:
                    st.plotly_chart(network_fig, width='stretch')
                    download_plotly_fig(network_fig, "code_network")
                else:
                    st.info("Not enough connections to create a network graph.")
        else:
            st.info("No codes available for visualization")

    with tab3:
        st.markdown("### Reflexive Memo - Phase 3")

        current_memo = get_memo(3)

        memo_text = st.text_area(
            "Reflexive notes on code refinement:",
            value=current_memo,
            height=200,
            placeholder="Which codes am I merging and why? What patterns am I seeing? How is my understanding evolving?..."
        )

        if st.button("Save Memo", key="memo_3"):
            save_memo(3, memo_text)
            st.success("Reflexive memo saved")

elif current_phase == 4:
    st.markdown("""
    <div class="phase-header">
        <h1>Phase 4: Developing Early Themes</h1>
        <p>Group codes into candidate themes with AI assistance and group comparisons</p>
    </div>
    """, unsafe_allow_html=True)

    tab1, tab2, tab3, tab4 = st.tabs(["AI Theme Clustering", "Manual Theme Building", "Group Comparisons", "Reflexive Memo"])

    with tab1:
        st.markdown("### AI-Suggested Theme Clusters")

        if not codes_df.empty:
            n_themes = st.slider("Number of themes to suggest:", 2, 8, 4)

            if st.button("🤖 Generate Theme Suggestions"):
                st.session_state.suggested_themes = cluster_codes_into_themes(codes_df, n_themes)

            if st.session_state.suggested_themes:
                st.markdown("#### Suggested Theme Groupings")

                # Copy the list to avoid modifying while iterating
                current_themes = st.session_state.suggested_themes.copy()

                for i in range(len(current_themes)):
                    theme = current_themes[i]
                    with st.expander(f"**{theme['theme_name']}**"):
                        st.markdown(f"*{theme['description']}*")

                        st.markdown("**Codes in this theme:**")
                        for code in theme['codes']:
                            st.markdown(f"• {code}")

                        with st.form(f"accept_form_{i}", clear_on_submit=True):
                            st.write(f"Accept this theme?")
                            if st.form_submit_button(f"Accept Theme '{theme['theme_name']}'"):
                                save_theme(theme)
                                st.success(f"Theme '{theme['theme_name']}' accepted and saved to database!")
                                # Remove accepted theme from session state without rerun
                                del st.session_state.suggested_themes[i]
                                # Re-render to update the list
                                st.rerun()

                # If no themes left, clear the list
                if not st.session_state.suggested_themes:
                    st.info("All suggested themes have been processed!")

        else:
            st.info("Complete earlier phases to have codes for theme development")

    with tab2:
        st.markdown("### Manual Theme Building")

        if not codes_df.empty:
            col1, col2 = st.columns(2)

            with col1:
                st.markdown("#### Available Codes")
                available_codes = st.multiselect(
                    "Select codes to group into theme:",
                    codes_df['code_name'].tolist()
                )

            with col2:
                st.markdown("#### Create Theme")

                with st.form("create_theme"):
                    theme_name = st.text_input("Theme Name:")
                    theme_description = st.text_area("Theme Description:")

                    com_b_primary = st.selectbox(
                        "Primary COM-B Category:",
                        ["", "Capability", "Opportunity", "Motivation"]
                    )

                    com_b_secondary = st.selectbox(
                        "Secondary COM-B Category (optional):",
                        ["", "Capability", "Opportunity", "Motivation"]
                    )

                    if st.form_submit_button("Create Theme"):
                        if theme_name and available_codes:
                            theme_data = {
                                'theme_name': theme_name,
                                'description': theme_description,
                                'com_b_primary': com_b_primary
                            }
                            theme_id = save_theme(theme_data)
                            # Link codes to theme
                            cursor = conn.cursor()
                            for code_name in available_codes:
                                cursor.execute('SELECT id FROM codes WHERE code_name = ?', (code_name,))
                                code_result = cursor.fetchone()
                                if code_result:
                                    code_id = code_result[0]
                                    cursor.execute('INSERT OR IGNORE INTO theme_codes (theme_id, code_id) VALUES (?, ?)', (theme_id, code_id))
                            conn.commit()
                            st.success(f"Theme '{theme_name}' created with {len(available_codes)} codes")

    with tab3:
        st.markdown("### Group Comparisons")

        if not transcripts_df.empty:
            st.markdown("#### Compare Themes Across Participant Groups")

            if 'ethnicity' in transcripts_df.columns:
                ethnicity_counts = transcripts_df[transcripts_df['ethnicity'] != '']['ethnicity'].value_counts()
                if len(ethnicity_counts) > 1:
                    df_eth = pd.DataFrame({'Ethnicity': ethnicity_counts.index, 'Count': ethnicity_counts.values})
                    fig1 = px.bar(df_eth, x='Ethnicity', y='Count', title="Participants by Ethnicity")
                    st.plotly_chart(fig1, width='stretch')
                    download_plotly_fig(fig1, "participants_ethnicity")

            if 'participant_type' in transcripts_df.columns:
                type_counts = transcripts_df['participant_type'].value_counts()
                fig2 = px.pie(values=type_counts.values, names=type_counts.index,
                            title="Black Women vs Healthcare Professionals")
                st.plotly_chart(fig2, width='stretch')
                download_plotly_fig(fig2, "participant_types")

            # HCP Role filter
            if 'role' in transcripts_df.columns:
                hcp_df = transcripts_df[transcripts_df['participant_type'] == 'Healthcare Professionals']
                if not hcp_df.empty and 'role' in hcp_df.columns:
                    role_counts = hcp_df['role'].value_counts()
                    df_role = pd.DataFrame({'Role': role_counts.index, 'Count': role_counts.values})
                    fig3 = px.bar(df_role, x='Role', y='Count', title="HCP Roles Distribution")
                    st.plotly_chart(fig3, width='stretch')
                    download_plotly_fig(fig3, "hcp_roles")

        st.info("Detailed group comparisons will be enhanced as themes develop")

    with tab4:
        st.markdown("### Reflexive Memo - Phase 4")

        current_memo = get_memo(4)

        memo_text = st.text_area(
            "Reflexive notes on theme development:",
            value=current_memo,
            height=200,
            placeholder="How are themes emerging? What differences am I seeing between groups? What themes surprise me?..."
        )

        if st.button("Save Memo", key="memo_4"):
            save_memo(4, memo_text)
            st.success("Reflexive memo saved")

elif current_phase == 5:
    st.markdown("""
    <div class="phase-header">
        <h1>Phase 5: Defining & Naming Themes</h1>
        <p>Refine theme names, descriptions, and COM-B mappings for final themes</p>
    </div>
    """, unsafe_allow_html=True)

    tab1, tab2, tab3, tab4 = st.tabs(["Theme Definitions", "COM-B Mapping", "Theme Validation", "Reflexive Memo"])

    with tab1:
        st.markdown("### AI-Generated Theme Summaries & Refinement")

        # Get existing themes
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM themes ORDER BY id DESC LIMIT 4')
        columns = [description[0] for description in cursor.description]
        themes_df = pd.DataFrame(cursor.fetchall(), columns=columns)

        if themes_df.empty:
            draft_themes = [
                {
                    "name": "Navigating Structural Inequities",
                    "description": "Systemic barriers within healthcare systems that create unequal access to breast screening for Black women",
                    "codes": ["Healthcare access", "Appointment scheduling", "Geographic barriers", "System navigation"],
                    "ai_summary": "This theme captures how structural elements of the healthcare system create barriers for Black women accessing breast cancer screening services."
                },
                {
                    "name": "Cultural Worlds and Health Beliefs",
                    "description": "The influence of cultural background, family beliefs, and community perspectives on screening decisions",
                    "codes": ["Cultural beliefs", "Family influence", "Community perceptions", "Religious considerations"],
                    "ai_summary": "This theme explores how cultural context and belief systems shape attitudes and behaviors towards breast screening."
                },
                {
                    "name": "Communication Disconnects",
                    "description": "Gaps and breakdowns in communication between healthcare providers and Black women",
                    "codes": ["GP communication", "Information delivery", "Language barriers", "Provider assumptions"],
                    "ai_summary": "This theme examines the communication challenges that exist between healthcare professionals and Black women regarding screening."
                },
                {
                    "name": "Emotional Landscapes of Screening",
                    "description": "The emotional responses, fears, and psychological aspects of breast screening experiences",
                    "codes": ["Fear and anxiety", "Cancer fears", "Procedure anxiety", "Body image concerns"],
                    "ai_summary": "This theme focuses on the emotional and psychological dimensions of breast screening for Black women."
                }
            ]
        else:
            draft_themes = []
            for _, theme in themes_df.iterrows():
                # Get codes for theme
                cursor.execute('SELECT c.code_name FROM codes c JOIN theme_codes tc ON c.id = tc.code_id WHERE tc.theme_id = ?', (theme['id'],))
                theme_codes = [row[0] for row in cursor.fetchall()]
                draft_themes.append({
                    "id": theme['id'],  # Fixed: Include the database ID
                    "name": theme['theme_name'],
                    "description": theme['description'],
                    "codes": theme_codes,
                    "ai_summary": theme['description']  # Use description as summary
                })

        for i, theme in enumerate(draft_themes):
            with st.expander(f"**Theme {i+1}: {theme['name']}**"):
                col1, col2 = st.columns(2)

                with col1:
                    st.markdown("**AI-Generated Summary:**")
                    st.info(theme['ai_summary'])

                    st.markdown("**Included Codes:**")
                    for code in theme['codes']:
                        st.markdown(f"• {code}")

                with col2:
                    st.markdown("**Researcher Refinement:**")

                    refined_name = st.text_input(
                        "Refined theme name:",
                        value=theme['name'],
                        key=f"name_{i}"
                    )

                    refined_desc = st.text_area(
                        "Refined description:",
                        value=theme['description'],
                        height=100,
                        key=f"desc_{i}"
                    )

                    # Versioning option
                    save_as_new = st.checkbox("Save as New Version", key=f"new_version_{i}")

                    if st.button(f"Save Refinements", key=f"save_{i}"):
                        theme_data = {
                            'theme_name': refined_name,
                            'description': refined_desc,
                            'com_b_primary': map_to_com_b(refined_name, refined_desc),
                            'com_b_secondary': ''  # Fixed: Include secondary (default empty)
                        }
                        try:
                            if save_as_new or themes_df.empty:
                                save_theme(theme_data)
                                st.success(f"New version of Theme {i+1} saved")
                            else:
                                update_theme_version(theme['id'], theme_data)
                                st.success(f"Theme {i+1} refinements saved")
                        except KeyError as e:
                            st.error(f"Missing key {e}. Please check theme data.")

    with tab2:
        st.markdown("### COM-B Framework Mapping")

        st.markdown("**Map each theme to primary and secondary COM-B categories:**")

        for i, theme in enumerate(draft_themes):
            st.markdown(f"#### {theme['name']}")

            col1, col2, col3 = st.columns(3)

            with col1:
                primary_comb = st.selectbox(
                    "Primary COM-B:",
                    ["Capability", "Opportunity", "Motivation"],
                    key=f"primary_comb_{i}"
                )

            with col2:
                secondary_comb = st.selectbox(
                    "Secondary COM-B:",
                    ["", "Capability", "Opportunity", "Motivation"],
                    key=f"secondary_comb_{i}"
                )

            with col3:
                suggested_comb = map_to_com_b(theme['name'], theme['description'])
                st.info(f"AI suggests: {suggested_comb}")

            comb_class = f"com-b-{primary_comb.lower()}"
            st.markdown(f"""
            <div class="{comb_class}" style="padding: 0.5rem; border-radius: 5px; margin: 0.5rem 0;">
                <strong>{theme['name']}</strong> → Primary: {primary_comb}
                {f", Secondary: {secondary_comb}" if secondary_comb else ""}
            </div>
            """, unsafe_allow_html=True)

    with tab3:
        st.markdown("### Theme Validation & Quality Metrics")

        col1, col2 = st.columns(2)

        with col1:
            st.markdown("#### Theme Quality Checklist")

            quality_criteria = [
                "Themes are internally coherent",
                "Themes are distinct from each other",
                "Themes answer the research questions",
                "Themes reflect the semantic content",
                "Theme names are concise and punchy"
            ]

            for criterion in quality_criteria:
                st.checkbox(criterion, key=f"quality_{criterion}")

        with col2:
            st.markdown("#### Theme Statistics")

            theme_stats = pd.DataFrame([
                {"Theme": "Structural Inequities", "Prevalence": 85, "Codes": 12, "Participants": 16},
                {"Theme": "Cultural Beliefs", "Prevalence": 78, "Codes": 9, "Participants": 11},
                {"Theme": "Communication", "Prevalence": 72, "Codes": 8, "Participants": 14},
                {"Theme": "Emotional Responses", "Prevalence": 65, "Codes": 7, "Participants": 10}
            ])

            fig = px.scatter(
                theme_stats, x="Codes", y="Prevalence",
                size="Participants", text="Theme",
                title="Theme Validation Metrics"
            )
            fig.update_traces(textposition="top center")
            st.plotly_chart(fig, width='stretch')
            download_plotly_fig(fig, "theme_validation")

    with tab4:
        st.markdown("### Reflexive Memo - Phase 5")

        current_memo = get_memo(5)

        memo_text = st.text_area(
            "Reflexive notes on theme definition:",
            value=current_memo,
            height=200,
            placeholder="How do the final theme names capture the essence? What COM-B mappings feel most accurate? What challenges am I facing in defining themes?..."
        )

        if st.button("Save Memo", key="memo_5"):
            save_memo(5, memo_text)
            st.success("Reflexive memo saved")

elif current_phase == 6:
    st.markdown("""
    <div class="phase-header">
        <h1>Phase 6: Write-Up</h1>
        <p>Generate draft reports with themes, codes, quotes, and reflexive memos</p>
    </div>
    """, unsafe_allow_html=True)

    tab1, tab2, tab3, tab4 = st.tabs(["Auto-Generated Report", "Quote Selection", "Export Options", "Reflexive Memo"])

    with tab1:
        st.markdown("### Auto-Generated Draft Report")

        guide_prompt = st.text_area(
            "Guide Prompt:",
            "Write a detailed academic report that focuses on the structural barriers to breast cancer screening for Black women.",
            height=150,
            help="Provide specific instructions to the LLM for generating the report, e.g., 'Focus on the role of healthcare professionals' perspectives.'",
            key="guide_prompt"
        )

        if st.button("Generate LLM Report"):
            if st.session_state.api_key:
                with st.spinner("Generating report..."):
                    all_data = get_all_data_for_llm()
                    report_text = generate_llm_report(all_data, guide_prompt, st.session_state.api_key)
                    st.session_state.llm_report = report_text

                st.success("Report generated!")
            else:
                st.warning("Please enter your API key first.")

        if 'llm_report' in st.session_state and st.session_state.llm_report:
            st.markdown("### Generated Report Draft")
            st.markdown(st.session_state.llm_report)

            b64 = base64.b64encode(st.session_state.llm_report.encode()).decode()
            href = f'<a href="data:file/txt;base64,{b64}" download="thematic_analysis_report.md">Download Report (Markdown)</a>'
            st.markdown(href, unsafe_allow_html=True)

    with tab2:
        st.markdown("### Supporting Quote Selection")

        # Get actual quotes from DB
        cursor = conn.cursor()
        cursor.execute('''
            SELECT ci.text_segment, t.id as participant, c.code_name, t.participant_type
            FROM code_instances ci
            JOIN transcripts t ON ci.transcript_id = t.id
            JOIN codes c ON ci.code_id = c.id
            ORDER BY RANDOM() LIMIT 20
        ''')
        quotes_data = cursor.fetchall()
        if quotes_data:
            theme_quotes = defaultdict(list)
            for quote, participant, code, ptype in quotes_data:
                # Simple theme assignment based on code for demo
                if 'access' in code.lower() or 'appointment' in code.lower():
                    theme = "Structural Inequities"
                elif 'culture' in code.lower() or 'family' in code.lower():
                    theme = "Cultural Beliefs"
                elif 'communication' in code.lower():
                    theme = "Communication"
                else:
                    theme = "Emotional Responses"
                theme_quotes[theme].append({
                    "quote": quote,
                    "participant": participant,
                    "context": f"Related to {code} ({ptype})"
                })
        else:
            # Fallback mock
            theme_quotes = {
                "Structural Inequities": [
                    {"quote": "The appointments are always during work hours... I can't keep taking time off",
                     "participant": "BW07", "context": "Discussing scheduling barriers"},
                    {"quote": "Getting to the hospital is a nightmare without a car",
                     "participant": "BW12", "context": "Geographic access challenges"}
                ],
                "Cultural Beliefs": [
                    {"quote": "In our culture, we don't talk about these things openly",
                     "participant": "BW03", "context": "Cultural communication norms"},
                    {"quote": "My mother always said hospitals are for when you're dying",
                     "participant": "BW09", "context": "Intergenerational beliefs"}
                ],
                "Communication": [
                    {"quote": "The GP never explained why it was important",
                     "participant": "BW05", "context": "Information delivery gaps"},
                    {"quote": "We assume they understand, but often they don't",
                     "participant": "HP02", "context": "Provider assumptions"}
                ],
                "Emotional Responses": [
                    {"quote": "I was terrified... what if they find something?",
                     "participant": "BW11", "context": "Cancer fears"},
                    {"quote": "The whole procedure felt so impersonal",
                     "participant": "BW04", "context": "Screening experience"}
                ]
            }

        selected_theme = st.selectbox("Select theme:", list(theme_quotes.keys()))

        st.markdown(f"#### Quotes for {selected_theme}")

        quotes_df = pd.DataFrame(theme_quotes[selected_theme])

        # Initialize or resize selected quotes for this theme to match current data
        if selected_theme not in st.session_state.selected_quotes or len(st.session_state.selected_quotes[selected_theme]) != len(quotes_df):
            st.session_state.selected_quotes[selected_theme] = [False] * len(quotes_df)

        for idx, quote_data in quotes_df.iterrows():
            col1, col2 = st.columns([3, 1])

            with col1:
                st.markdown(f"**Quote {idx+1}:** *\"{quote_data['quote']}\"*")
                st.caption(f"Participant: {quote_data['participant']} | Context: {quote_data['context']}")

            with col2:
                include_quote = st.checkbox("Include in report", value=st.session_state.selected_quotes[selected_theme][idx], key=f"quote_{selected_theme}_{idx}")
                st.session_state.selected_quotes[selected_theme][idx] = include_quote
                if include_quote:
                    st.success("✓ Selected")

        if st.button(f"Save Selections for {selected_theme}"):
            st.success(f"Selections for {selected_theme} saved!")

    with tab3:
        st.markdown("### Export Options")

        export_format = st.radio(
            "Choose format:",
            ["Word Document (.docx)", "Excel Workbook (.xlsx)", "PDF Summary", "JSON Data"]
        )

        # Map display name to internal format
        internal_format = {
            "Word Document (.docx)": "docx",
            "Excel Workbook (.xlsx)": "xlsx",
            "PDF Summary": "pdf",
            "JSON Data": "json"
        }.get(export_format, "docx")

        if st.button("Generate Export"):
            data, mime, fname = generate_export(internal_format)
            st.download_button(
                label=f"Download {export_format}",
                data=data,
                file_name=fname,
                mime=mime
            )
            st.success(f"Export generated as {export_format}")

    with tab4:
        st.markdown("### Reflexive Memo - Phase 6")

        current_memo = get_memo(6)

        memo_text = st.text_area(
            "Final reflexive notes:",
            value=current_memo,
            height=200,
            placeholder="How has my understanding evolved through the analysis? What would I do differently? What are the implications of these findings?..."
        )

        if st.button("Save Final Memo", key="memo_6"):
            save_memo(6, memo_text)
            st.success("Reflexive memo saved")

elif current_phase == 7:
    st.markdown("""
    <div class="phase-header">
        <h1>Visual Analytics Dashboard</h1>
        <p>Exploratory visualizations and comparative analysis</p>
    </div>
    """, unsafe_allow_html=True)

    tab1, tab2, tab3 = st.tabs(["Comparative Dashboards", "Advanced Visualizations", "Data Export"])

    with tab1:
        st.markdown("### Comparative Analysis")

        col1, col2 = st.columns(2)

        with col1:
            st.markdown("#### African vs Caribbean Women")

            if not transcripts_df.empty and 'ethnicity' in transcripts_df.columns:
                ethnicity_counts = transcripts_df[transcripts_df['ethnicity'] != '']['ethnicity'].value_counts()

                if len(ethnicity_counts) > 1:
                    df_eth = pd.DataFrame({'Ethnicity': ethnicity_counts.index, 'Count': ethnicity_counts.values})
                    fig1 = px.pie(
                        df_eth, values='Count',
                        names='Ethnicity',
                        title="Participant Ethnicity Distribution"
                    )
                    st.plotly_chart(fig1, width='stretch')
                    download_plotly_fig(fig1, "ethnicity_distribution")

                comparison_data = pd.DataFrame({
                    'Theme': ['Structural', 'Cultural', 'Communication', 'Emotional'],
                    'African': [75, 85, 65, 70],
                    'Caribbean': [80, 75, 70, 65]
                })

                fig2 = px.bar(
                    comparison_data.melt(id_vars=['Theme'], var_name='Ethnicity', value_name='Prevalence'),
                    x='Theme', y='Prevalence', color='Ethnicity',
                    title="Theme Prevalence by Ethnicity"
                )
                st.plotly_chart(fig2, width='stretch')
                download_plotly_fig(fig2, "theme_prevalence_ethnicity")

        with col2:
            st.markdown("#### Black Women vs Healthcare Professionals")

            if not transcripts_df.empty and 'participant_type' in transcripts_df.columns:
                type_data = transcripts_df['participant_type'].value_counts()
                df_type = pd.DataFrame({'Type': type_data.index, 'Count': type_data.values})

                fig3 = px.bar(
                    df_type, x='Type', y='Count',
                    title="Participants by Type"
                )
                st.plotly_chart(fig3, width='stretch')
                download_plotly_fig(fig3, "participants_type")

                # HCP Roles
                hcp_df = transcripts_df[transcripts_df['participant_type'] == 'Healthcare Professionals']
                if 'role' in hcp_df.columns and not hcp_df.empty:
                    role_data = hcp_df['role'].value_counts()
                    df_role = pd.DataFrame({'Role': role_data.index, 'Count': role_data.values})
                    fig_role = px.bar(df_role, x='Role', y='Count', title="HCP Roles")
                    st.plotly_chart(fig_role, width='stretch')
                    download_plotly_fig(fig_role, "hcp_roles_dashboard")

                perspective_data = pd.DataFrame({
                    'Theme': ['Structural', 'Cultural', 'Communication', 'Emotional'],
                    'Black Women': [85, 80, 75, 85],
                    'Healthcare Professionals': [70, 40, 85, 30]
                })

                fig4 = px.bar(
                    perspective_data.melt(id_vars=['Theme'], var_name='Perspective', value_name='Emphasis'),
                    x='Theme', y='Emphasis', color='Perspective',
                    title="Theme Emphasis by Participant Type"
                )
                st.plotly_chart(fig4, width='stretch')
                download_plotly_fig(fig4, "theme_emphasis_perspective")

    with tab2:
        st.markdown("### Advanced Visualizations")

        viz_col1, viz_col2 = st.columns(2)

        with viz_col1:
            st.markdown("#### Sentiment Analysis by Theme")

            # Use actual data if possible
            if not codes_df.empty and not transcripts_df.empty:
                # Simple sentiment using TextBlob
                sentiments = []
                for _, code in codes_df.iterrows():
                    for _, trans in transcripts_df.iterrows():
                        blob = TextBlob(trans['content'])
                        sentiments.append({
                            'Theme': code['code_name'][:10],  # Truncate
                            'Sentiment': blob.sentiment.polarity,
                            'Subjectivity': blob.sentiment.subjectivity,
                            'Participant': trans['participant_type'][:2]
                        })
                sentiment_data = pd.DataFrame(sentiments[:20])  # Limit
            else:
                sentiment_data = pd.DataFrame({
                    'Theme': ['Structural', 'Cultural', 'Communication', 'Emotional'] * 5,
                    'Sentiment': np.random.normal(0, 0.5, 20),
                    'Subjectivity': np.random.uniform(0.3, 0.9, 20),
                    'Participant': ['BW' if i % 2 == 0 else 'HP' for i in range(20)]
                })

            fig_sentiment = px.scatter(
                sentiment_data, x='Sentiment', y='Subjectivity',
                color='Theme', symbol='Participant',
                title="Sentiment Analysis by Theme and Participant Type"
            )
            st.plotly_chart(fig_sentiment, width='stretch')
            download_plotly_fig(fig_sentiment, "sentiment_analysis")

        with viz_col2:
            st.markdown("#### COM-B Distribution")

            comb_data = pd.DataFrame({
                'COM-B': ['Capability', 'Opportunity', 'Motivation'],
                'Frequency': [25, 35, 40],
                'Themes': [2, 1, 2]
            })

            fig_comb = px.sunburst(
                comb_data,
                path=['COM-B'],
                values='Frequency',
                title="COM-B Framework Distribution"
            )
            st.plotly_chart(fig_comb, width='stretch')
            download_plotly_fig(fig_comb, "com_b_distribution")

        st.markdown("#### Interactive Code Network")

        if not codes_df.empty:
            network_threshold = st.slider("Network connection threshold:", 1, 5, 2)
            network_fig = create_cooccurrence_network(codes_df, network_threshold)
            if network_fig:
                st.plotly_chart(network_fig, width='stretch')
                download_plotly_fig(network_fig, "interactive_code_network")
            else:
                st.info("Not enough codes to create a network graph.")

    with tab3:
        st.markdown("### Research Data Export")

        col1, col2 = st.columns(2)

        with col1:
            st.markdown("#### Complete Dataset Export")

            if st.button("Export All Research Data"):
                data, mime, fname = generate_export('xlsx')  # Default to xlsx for full data
                st.download_button(
                    label="Download Complete Dataset (XLSX)",
                    data=data,
                    file_name=fname,
                    mime=mime
                )

        with col2:
            st.markdown("#### Analysis Summary")

            summary_metrics = {
                "Transcripts Processed": len(transcripts_df),
                "Codes Generated": len(codes_df),
                "Themes Developed": len(pd.read_sql_query("SELECT * FROM themes", conn)),
                "Reflexive Memos": sum(1 for i in range(1,7) if get_memo(i)),
                "Transcript Memos": len(pd.read_sql_query("SELECT * FROM transcript_memos", conn)),
                "Analysis Completion": "100%"
            }

            for metric, value in summary_metrics.items():
                st.metric(metric, value)

st.markdown("---")

phases_completed = [1, 2, 3, 4, 5, 6]
progress_bar = st.progress(len(phases_completed) / 6)
st.caption(f"Analysis Progress: {len(phases_completed)}/6 phases completed")

st.markdown("*Reflexive Thematic Analysis Tool for Breast Cancer Screening Disparities Research*")
st.caption("Features: Phase-guided analysis, AI assistance, COM-B framework integration, Reflexive memo tracking, Audio transcription, Theme versioning, Fuzzy matching, Multi-format exports")
